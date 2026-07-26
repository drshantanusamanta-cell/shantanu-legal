"""
Export drafts and opinions to DOCX, PDF and Markdown.

DOCX output uses a court-friendly setup: Times New Roman 12pt, 1.5 line
spacing, justified body, numbered paragraphs preserved, and an appended
verification appendix listing every citation with its verified source link.

The appendix is deliberate. A draft that leaves the app should carry proof of
what was checked and what was suppressed.
"""

from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Any

try:
    from config import COPYRIGHT_HOLDER, COPYRIGHT_NOTICE
except Exception:  # pragma: no cover
    COPYRIGHT_HOLDER = "Dr Shantanu Samanta"
    COPYRIGHT_NOTICE = "© 2026 Dr Shantanu Samanta. All rights reserved."


# --------------------------------------------------------------------------
# Markdown
# --------------------------------------------------------------------------
def to_markdown(title: str, body: str, report: Any | None = None) -> bytes:
    parts = [f"# {title}", "", body, ""]
    if report is not None:
        parts.append(_verification_markdown(report))
    return "\n".join(parts).encode("utf-8")


def _verification_markdown(report: Any) -> str:
    lines = ["", "---", "", "## Verification Appendix", ""]
    lines.append(
        f"- Authorities asserted: **{report.total}**  "
        f"| verified: **{len(report.verified)}**  "
        f"| suppressed: **{len(report.suppressed)}**"
    )
    lines.append(f"- Quoted passages checked: **{report.quotes_checked}**, "
                 f"failed: **{len(report.quotes_failed)}**")
    lines.append(f"- Sources consulted: {', '.join(report.sources_consulted) or 'none'}")
    lines.append("")
    if report.verified:
        lines.append("### Verified authorities")
        for c in report.verified:
            lines.append(
                f"- **{c.verified_title or c.title}** "
                f"{('— ' + c.reporter) if c.reporter else ''} "
                f"({c.verified_via}) {c.verified_url}"
            )
        lines.append("")
    if report.suppressed:
        lines.append("### Suppressed — COULD NOT BE VERIFIED")
        for c in report.suppressed:
            lines.append(f"- `{c.label()}` — {c.failure_reason}")
        lines.append("")
    lines.append(
        "> Generated with AI assistance. Every authority above must be independently "
        "confirmed against the certified report before filing."
    )
    lines.append("")
    lines.append("---")
    lines.append(f"*{COPYRIGHT_NOTICE}*")
    return "\n".join(lines)


# --------------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------------
def to_docx(title: str, body: str, report: Any | None = None,
            court_style: bool = True) -> bytes:
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    document = docx.Document()

    # Page + base style
    for section in document.sections:
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    h = document.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(14)

    for raw_line in (body or "").split("\n"):
        line = raw_line.rstrip()
        if not line.strip():
            document.add_paragraph()
            continue

        # Headings
        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            p = document.add_paragraph()
            r = p.add_run(_strip_md(m.group(2)))
            r.bold = True
            r.font.size = Pt(14 - level)
            p.paragraph_format.space_before = Pt(10)
            continue

        # Bullets
        if re.match(r"^\s*[-*•]\s+", line):
            p = document.add_paragraph(style="List Bullet")
            _add_rich(p, re.sub(r"^\s*[-*•]\s+", "", line))
            continue

        # Numbered paragraphs: keep the number literal (court drafts rely on it)
        p = document.add_paragraph()
        if court_style:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_rich(p, line)

    if report is not None:
        document.add_page_break()
        ph = document.add_paragraph()
        r = ph.add_run("VERIFICATION APPENDIX")
        r.bold = True
        r.font.size = Pt(13)

        document.add_paragraph(
            f"Authorities asserted: {report.total}  |  Verified: {len(report.verified)}  "
            f"|  Suppressed: {len(report.suppressed)}"
        )
        document.add_paragraph(
            f"Quoted passages checked: {report.quotes_checked}  |  "
            f"Failed verification: {len(report.quotes_failed)}"
        )
        document.add_paragraph(
            f"Sources consulted: {', '.join(report.sources_consulted) or 'none'}"
        )
        document.add_paragraph(
            f"Generated: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
        )

        if report.verified:
            sp = document.add_paragraph()
            sp.add_run("Verified authorities").bold = True
            for c in report.verified:
                p = document.add_paragraph(style="List Bullet")
                p.add_run(c.verified_title or c.title).bold = True
                tail = f"  {c.reporter}  [{c.verified_via}]  {c.verified_url}"
                p.add_run(tail).font.size = Pt(10)

        if report.suppressed:
            sp = document.add_paragraph()
            rr = sp.add_run("Suppressed — could not be verified")
            rr.bold = True
            rr.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            for c in report.suppressed:
                p = document.add_paragraph(style="List Bullet")
                run = p.add_run(f"{c.label()} — {c.failure_reason}")
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

        note = document.add_paragraph()
        nr = note.add_run(
            "Prepared with AI assistance. Every authority must be independently "
            "confirmed against the certified report before filing or service."
        )
        nr.italic = True
        nr.font.size = Pt(9)

    # Attribution footer on every exported document.
    cp = document.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run(COPYRIGHT_NOTICE)
    cr.font.size = Pt(8)
    cr.italic = True
    cr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Repeat it in the running page footer so it survives extraction.
    try:
        for section in document.sections:
            footer_p = section.footer.paragraphs[0]
            footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fr = footer_p.add_run(COPYRIGHT_NOTICE)
            fr.font.size = Pt(7.5)
            fr.italic = True
            fr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    except Exception:
        pass

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _strip_md(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\((.+?)\)", r"\1 (\2)", text)
    return text


def _add_rich(paragraph, text: str) -> None:
    """Render **bold** / *italic* / [link](url) inline into a docx paragraph."""
    text = re.sub(r"\[(.+?)\]\((.+?)\)", r"\1 (\2)", text)
    tokens = re.split(r"(\*\*.+?\*\*|\*[^*]+?\*|`.+?`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            paragraph.add_run(tok[2:-2]).bold = True
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            paragraph.add_run(tok[1:-1]).italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = paragraph.add_run(tok[1:-1])
            r.font.name = "Courier New"
        else:
            paragraph.add_run(tok)


# --------------------------------------------------------------------------
# PDF
# --------------------------------------------------------------------------
def to_pdf(title: str, body: str, report: Any | None = None) -> bytes:
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=1.4 * inch, rightMargin=1.0 * inch,
        topMargin=1.0 * inch, bottomMargin=1.0 * inch,
        title=title, author=COPYRIGHT_HOLDER, creator=COPYRIGHT_HOLDER,
        subject=COPYRIGHT_NOTICE,
    )

    def _stamp(canvas, doc_):
        """Copyright + page number on every page."""
        canvas.saveState()
        canvas.setFont("Times-Italic", 7.5)
        canvas.setFillColorRGB(0.58, 0.64, 0.72)
        canvas.drawCentredString(A4[0] / 2.0, 0.55 * inch, COPYRIGHT_NOTICE)
        canvas.setFont("Times-Roman", 8)
        canvas.drawRightString(A4[0] - 1.0 * inch, 0.55 * inch, f"Page {doc_.page}")
        canvas.restoreState()

    ss = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "LegalBody", parent=ss["Normal"], fontName="Times-Roman", fontSize=11.5,
        leading=19, alignment=TA_JUSTIFY, spaceAfter=8,
    )
    title_style = ParagraphStyle(
        "LegalTitle", parent=ss["Title"], fontName="Times-Bold", fontSize=15,
        alignment=TA_CENTER, spaceAfter=18,
    )
    head_style = ParagraphStyle(
        "LegalHead", parent=ss["Heading2"], fontName="Times-Bold", fontSize=13,
        spaceBefore=12, spaceAfter=6,
    )
    small = ParagraphStyle(
        "Small", parent=ss["Normal"], fontName="Times-Roman", fontSize=9, leading=12
    )

    def esc(t: str) -> str:
        t = t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        t = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", t)
        t = re.sub(r"(?<!\*)\*([^*]+?)\*(?!\*)", r"<i>\1</i>", t)
        t = re.sub(r"\[(.+?)\]\((.+?)\)", r'<link href="\2"><u>\1</u></link>', t)
        return t

    flow: list[Any] = [Paragraph(esc(title.upper()), title_style)]

    for raw in (body or "").split("\n"):
        line = raw.rstrip()
        if not line.strip():
            flow.append(Spacer(1, 6))
            continue
        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            flow.append(Paragraph(esc(m.group(2)), head_style))
            continue
        if re.match(r"^\s*[-*•]\s+", line):
            txt = re.sub(r"^\s*[-*•]\s+", "• ", line)
            flow.append(Paragraph(esc(txt), body_style))
            continue
        flow.append(Paragraph(esc(line), body_style))

    if report is not None:
        flow.append(PageBreak())
        flow.append(Paragraph("VERIFICATION APPENDIX", title_style))
        flow.append(Paragraph(
            f"Authorities asserted: <b>{report.total}</b> &nbsp;|&nbsp; "
            f"Verified: <b>{len(report.verified)}</b> &nbsp;|&nbsp; "
            f"Suppressed: <b>{len(report.suppressed)}</b>", body_style))
        flow.append(Paragraph(
            f"Quoted passages checked: {report.quotes_checked}, failed: "
            f"{len(report.quotes_failed)}", body_style))
        flow.append(Paragraph(
            f"Sources consulted: {', '.join(report.sources_consulted) or 'none'}",
            body_style))

        if report.verified:
            flow.append(Paragraph("Verified authorities", head_style))
            for c in report.verified:
                flow.append(Paragraph(
                    esc(f"• <b>{c.verified_title or c.title}</b> {c.reporter} "
                        f"[{c.verified_via}] {c.verified_url}"), small))
        if report.suppressed:
            flow.append(Paragraph("Suppressed — could not be verified", head_style))
            for c in report.suppressed:
                flow.append(Paragraph(
                    esc(f"• {c.label()} — {c.failure_reason}"), small))

        flow.append(Spacer(1, 14))
        flow.append(Paragraph(
            "<i>Prepared with AI assistance. Every authority must be independently "
            "confirmed against the certified report before filing or service.</i>",
            small))

    doc.build(flow, onFirstPage=_stamp, onLaterPages=_stamp)
    return buf.getvalue()


def safe_filename(text: str, ext: str) -> str:
    base = re.sub(r"[^\w\s\-]", "", text or "document").strip()
    base = re.sub(r"\s+", "_", base)[:60] or "document"
    stamp = datetime.now().strftime("%Y%m%d_%H%M")
    return f"{base}_{stamp}.{ext}"
