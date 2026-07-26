"""
Document ingestion: PDF, DOCX, images, plain text.

Strategy for PDFs, in order:
  1. Extract the embedded text layer (fast, free, exact).
  2. If that yields almost nothing, the PDF is a scan -> send the pages to the
     VISION model as images. Modern multimodal models beat Tesseract badly on
     Indian court documents (stamps, seals, handwriting, Devanagari, poor scans),
     so we prefer that over a local OCR dependency.
  3. Optional local OCR via pytesseract if installed and preferred offline.
"""

from __future__ import annotations

import base64
import io
from dataclasses import dataclass, field
from typing import Any

from llm.base import Attachment


MIN_CHARS_PER_PAGE_FOR_TEXT_LAYER = 80
MAX_VISION_PAGES = 20


@dataclass
class IngestedDocument:
    filename: str
    media_type: str
    text: str = ""
    page_count: int = 0
    needs_vision: bool = False
    attachments: list[Attachment] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    meta: dict[str, Any] = field(default_factory=dict)

    @property
    def has_text(self) -> bool:
        return len(self.text.strip()) > 40

    def preview(self, n: int = 1500) -> str:
        t = self.text.strip()
        return t[:n] + ("…" if len(t) > n else "")


def _b64(data: bytes) -> str:
    return base64.standard_b64encode(data).decode("utf-8")


# --------------------------------------------------------------------------
def ingest_pdf(data: bytes, filename: str, force_vision: bool = False) -> IngestedDocument:
    doc = IngestedDocument(filename=filename, media_type="application/pdf")

    text_parts: list[str] = []
    page_count = 0

    if not force_vision:
        try:
            import pdfplumber

            with pdfplumber.open(io.BytesIO(data)) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    try:
                        text_parts.append(page.extract_text() or "")
                    except Exception:
                        text_parts.append("")
        except Exception as exc:  # noqa: BLE001
            doc.warnings.append(f"pdfplumber failed ({exc}); trying pypdf.")
            try:
                from pypdf import PdfReader

                reader = PdfReader(io.BytesIO(data))
                page_count = len(reader.pages)
                for page in reader.pages:
                    try:
                        text_parts.append(page.extract_text() or "")
                    except Exception:
                        text_parts.append("")
            except Exception as exc2:  # noqa: BLE001
                doc.warnings.append(f"pypdf also failed ({exc2}).")

    doc.page_count = page_count
    doc.text = "\n\n".join(t for t in text_parts if t).strip()

    density = len(doc.text) / page_count if page_count else 0
    if force_vision or density < MIN_CHARS_PER_PAGE_FOR_TEXT_LAYER:
        doc.needs_vision = True
        if not force_vision:
            doc.warnings.append(
                f"Sparse text layer ({density:.0f} chars/page) — treating as a scanned "
                "document and sending it to the vision model."
            )
        # Send the whole PDF: Anthropic ingests PDFs natively, Gemini too.
        doc.attachments.append(
            Attachment(data_b64=_b64(data), media_type="application/pdf", filename=filename)
        )
        if page_count > MAX_VISION_PAGES:
            doc.warnings.append(
                f"{page_count} pages is large; consider splitting to control token cost."
            )
    return doc


def ingest_image(data: bytes, filename: str, media_type: str) -> IngestedDocument:
    doc = IngestedDocument(filename=filename, media_type=media_type, needs_vision=True)
    doc.attachments.append(
        Attachment(data_b64=_b64(data), media_type=media_type, filename=filename)
    )

    # Optional local OCR, purely as a bonus text layer.
    try:
        import pytesseract
        from PIL import Image

        img = Image.open(io.BytesIO(data))
        doc.text = pytesseract.image_to_string(img) or ""
    except Exception:
        pass
    return doc


def ingest_docx(data: bytes, filename: str) -> IngestedDocument:
    doc = IngestedDocument(
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    try:
        import docx

        d = docx.Document(io.BytesIO(data))
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        doc.text = "\n".join(parts).strip()
    except Exception as exc:  # noqa: BLE001
        doc.warnings.append(f"Could not read .docx: {exc}")
    return doc


def ingest_text(data: bytes, filename: str) -> IngestedDocument:
    doc = IngestedDocument(filename=filename, media_type="text/plain")
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            doc.text = data.decode(enc)
            break
        except Exception:
            continue
    return doc


EXT_MAP = {
    "pdf": "application/pdf",
    "png": "image/png",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "webp": "image/webp",
    "gif": "image/gif",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "txt": "text/plain",
    "md": "text/plain",
    "rtf": "text/plain",
}


def ingest(data: bytes, filename: str, force_vision: bool = False) -> IngestedDocument:
    """Dispatch on file extension."""
    ext = (filename.rsplit(".", 1)[-1] if "." in filename else "").lower()
    media = EXT_MAP.get(ext, "application/octet-stream")

    if media == "application/pdf":
        return ingest_pdf(data, filename, force_vision=force_vision)
    if media.startswith("image/"):
        return ingest_image(data, filename, media)
    if ext == "docx":
        return ingest_docx(data, filename)
    if media == "text/plain":
        return ingest_text(data, filename)

    doc = IngestedDocument(filename=filename, media_type=media)
    doc.warnings.append(f"Unsupported file type '.{ext}'. Supported: {', '.join(EXT_MAP)}.")
    return doc


def ingest_many(files: list[tuple[bytes, str]], force_vision: bool = False) -> list[IngestedDocument]:
    return [ingest(data, name, force_vision=force_vision) for data, name in files]


def combined_context(docs: list[IngestedDocument], max_chars: int = 120_000) -> str:
    """Flatten extracted text into a prompt block, with a budget."""
    out: list[str] = []
    used = 0
    for d in docs:
        if not d.has_text:
            continue
        header = f"\n\n===== DOCUMENT: {d.filename} ({d.page_count or '?'} pages) =====\n"
        body = d.text
        room = max_chars - used - len(header)
        if room <= 0:
            out.append("\n\n[Further documents omitted — context budget reached.]")
            break
        if len(body) > room:
            body = body[:room] + "\n[…truncated…]"
        out.append(header + body)
        used += len(header) + len(body)
    return "".join(out)


def all_attachments(docs: list[IngestedDocument]) -> list[Attachment]:
    atts: list[Attachment] = []
    for d in docs:
        atts.extend(d.attachments)
    return atts
